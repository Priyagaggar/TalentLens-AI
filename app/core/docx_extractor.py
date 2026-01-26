
import logging
import os
from typing import Optional

import docx
from docx.opc.exceptions import PackageNotFoundError

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocxExtractionError(Exception):
    """Custom exception for DOCX extraction failures."""
    pass


def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file, including paragraphs and tables.

    Args:
        file_path (str): The absolute path to the DOCX file.

    Returns:
        str: The extracted text as a single string.

    Raises:
        FileNotFoundError: If the file does not exist.
        DocxExtractionError: If file is not a valid DOCX or other IO errors.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    logger.info(f"Attempting to extract text from {file_path} using python-docx...")

    try:
        doc = docx.Document(file_path)
        full_text = []

        # 1. Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # 2. Extract from tables (common in resumes for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Cell content usually has paragraphs too
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

        extracted_text = "\n".join(full_text)
        
        if extracted_text.strip():
            logger.info("Successfully extracted text from DOCX.")
            return extracted_text.strip()
        else:
            logger.warning("DOCX file was processed but yielded no text.")
            return ""

    except PackageNotFoundError:
        logger.error(f"File {file_path} is not a valid DOCX file (PackageNotFoundError).")
        raise DocxExtractionError(f"Invalid DOCX file: {file_path}")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise DocxExtractionError(f"Failed to process DOCX file: {e}")


if __name__ == "__main__":
    # Test execution
    test_docx_path = "test_resume.docx"

    # Create dummy DOCX if it doesn't exist
    if not os.path.exists(test_docx_path):
        doc = docx.Document()
        doc.add_heading('John Doe', 0)
        doc.add_paragraph('Python Developer')
        
        # Add a table to test table extraction
        table = doc.add_table(rows=2, cols=2)
        cell1 = table.cell(0, 0)
        cell1.text = "Skills"
        cell2 = table.cell(0, 1)
        cell2.text = "Python, FastAPI, AI"
        
        cell3 = table.cell(1, 0)
        cell3.text = "Education"
        cell4 = table.cell(1, 1)
        cell4.text = "Computer Science"
        
        doc.save(test_docx_path)
        print(f"Created temporary test file: {test_docx_path}")

    try:
        extracted = extract_text_from_docx(test_docx_path)
        print("\n--- Extracted Text ---")
        print(extracted)
        print("----------------------\n")
    except Exception as e:
        print(f"Error: {e}")

    # cleanup
    if os.path.exists(test_docx_path):
        os.remove(test_docx_path)
