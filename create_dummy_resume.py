try:
    from docx import Document
    document = Document()
    document.add_heading('John Doe', 0)
    document.add_paragraph('Experience: 5 years in Python and React.')
    document.add_paragraph('Skills: Python, React, JavaScript, SQL.')
    document.add_paragraph('Education: BS in Computer Science.')
    document.save('d:/coding/ai_resume_screener/dummy_resume.docx')
    print("Successfully created dummy_resume.docx")
except ImportError:
    print("python-docx not found")
except Exception as e:
    print(f"Error: {e}")
